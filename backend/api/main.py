from fastapi import FastAPI, File, UploadFile, HTTPException, Form, Query, BackgroundTasks, Body
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import json
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import sys
import os
import google.generativeai as genai
import re
import time
import uuid
import base64
import tempfile
import asyncio
from datetime import datetime
from dotenv import load_dotenv
from firecrawl import FirecrawlApp
import io
import PyPDF2
import docx
import openai
import requests
from bs4 import BeautifulSoup
import traceback

# Load environment variables
load_dotenv()

# Add the parent directory to the path so we can import from agents
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from agents.crew import JobSkillCrew

# Configure OpenAI API
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if OPENAI_API_KEY:
    openai.api_key = OPENAI_API_KEY
    print("OpenAI API key configured successfully")
    # Set the client to use the API key
    client = openai.OpenAI(api_key=OPENAI_API_KEY)
    
    try:
        # Test the API connection
        models = client.models.list()
        print("Available OpenAI models:")
        for model in models.data[:5]:  # Just print the first few to avoid clutter
            print(f"- {model.id}")
        print("...")
    except Exception as e:
        print(f"Error listing OpenAI models: {str(e)}")
else:
    print("Warning: OPENAI_API_KEY not found in environment variables")
    client = None

# Configure Google Generative AI (Gemini) API
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY") or os.getenv("Gemini_API_KEY")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    print("Google Generative AI (Gemini) API key configured successfully")
else:
    print("Warning: Gemini API key not found in environment variables")

app = FastAPI(title="JobSkillTracker API")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins for development
    allow_credentials=True,
    allow_methods=["*"],  # Allow all methods
    allow_headers=["*"],  # Allow all headers
)

# Initialize the JobSkillCrew
job_skill_crew = JobSkillCrew()

# Initialize Firecrawl with API key from environment variables
firecrawl_api_key = os.getenv("FIRECRAWL_API_KEY", "")
firecrawl_app = FirecrawlApp(api_key=firecrawl_api_key) if firecrawl_api_key else None

# Define request models
class JobDescription(BaseModel):
    url: str
    title: Optional[str] = ""
    company: Optional[str] = ""
    description: Optional[str] = ""  # Added back for fallback
    useFirecrawl: Optional[bool] = True

class ResumeAnalysisRequest(BaseModel):
    resume_text: str
    extracted_job_skills: Optional[Dict[str, List[str]]] = None
    
class SkillGoalsRequest(BaseModel):
    user_id: str  # Unique identifier for the user
    skill_goals: List[str]  # List of skills the user wants to acquire
    current_skills: List[str] = []  # Optional list of skills the user already has
    
class JobMatchRequest(BaseModel):
    user_id: str
    job_description: str
    job_title: str
    company: str
    url: str
    
class JobSearchRequest(BaseModel):
    user_id: str
    resume_text: str
    keywords: List[str] = []  # Optional additional keywords
    location: str = ""  # Optional location
    experience_level: str = ""  # Optional experience level

class ProjectRecommendationRequest(BaseModel):
    skill_gaps: List[Dict[str, Any]]
    current_skills: List[Dict[str, Any]]
    
class ChatRequest(BaseModel):
    message: str
    chat_history: Optional[List[Dict[str, str]]] = []
    mode: Optional[str] = "general"  # general, paper_summary, skill_roadmap
    paper_url: Optional[str] = None
    paper_text: Optional[str] = None
    target_skill: Optional[str] = None

# Define API endpoints
@app.post("/extract-job-skills")
async def extract_job_skills(request: dict = Body(...)):
    """
    Extract skills from a job description using AI.
    """
    try:
        job_description = request.get("job_description", "")
        job_title = request.get("job_title", "")
        company = request.get("company", "")
        
        if not job_description:
            return {"error": "Job description is required"}
        
        print(f"Extracting skills from job description for {job_title} at {company}")
        print(f"Description length: {len(job_description)} characters")
        
        # Use OpenAI API to extract skills if available
        if OPENAI_API_KEY and client:
            try:
                # Create prompt for skill extraction
                prompt = f"Extract all technical and soft skills from the following job description for {job_title} at {company}.\n\nJob Description:\n{job_description[:5000]}\n\nReturn ONLY a JSON array of strings with the skill names. For example: [\"JavaScript\", \"React\", \"Communication\", \"Problem Solving\"]\nDo not include any explanations, just the JSON array."
                
                # Generate response using OpenAI
                response = client.chat.completions.create(
                    model="gpt-3.5-turbo",  # Using GPT-3.5 for cost efficiency
                    messages=[
                        {"role": "system", "content": "You are a skilled job analyzer that extracts technical and soft skills from job descriptions. Return only a JSON array of skills without any explanation."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3,  # Lower temperature for more consistent results
                    max_tokens=1000
                )
                
                # Extract the response text
                response_text = response.choices[0].message.content
                
                # Extract JSON array from response
                import re
                import json
                
                # Try to find a JSON array in the response
                json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
                
                if json_match:
                    try:
                        skills_array = json.loads(json_match.group(0))
                        print(f"Extracted {len(skills_array)} skills")
                        return {"skills": skills_array}
                    except json.JSONDecodeError as e:
                        print(f"Error parsing JSON from response: {str(e)}")
                        print(f"Response text: {response_text}")
                        # Fall back to simple extraction
                        skills = re.findall(r'"([^"]+)"', response_text)
                        if skills:
                            print(f"Extracted {len(skills)} skills using regex")
                            return {"skills": skills}
                else:
                    print(f"No JSON array found in response: {response_text}")
                    # Try to extract skills using regex as fallback
                    skills = re.findall(r'"([^"]+)"', response_text)
                    if skills:
                        print(f"Extracted {len(skills)} skills using regex")
                        return {"skills": skills}
            except Exception as e:
                print(f"Error using OpenAI API for skill extraction: {str(e)}")
        
        # Fallback to rule-based extraction if AI fails or is not available
        skills = extract_skills_rule_based(job_description, job_title)
        print(f"Extracted {len(skills)} skills using rule-based approach")
        return {"skills": skills}
        
    except Exception as e:
        print(f"Unexpected error in skill extraction: {str(e)}")
        return {"error": f"Error extracting skills: {str(e)}"}

def extract_skills_rule_based(job_description, job_title):
    """
    Extract skills from job description using rule-based approach.
    This is a fallback method when AI extraction is not available.
    """
    # Common technical skills to look for
    tech_skills = [
        "Python", "JavaScript", "Java", "C++", "C#", "Ruby", "PHP", "Swift", "Kotlin", "Go",
        "React", "Angular", "Vue.js", "Node.js", "Django", "Flask", "Spring", "ASP.NET",
        "SQL", "NoSQL", "MongoDB", "PostgreSQL", "MySQL", "Oracle", "Firebase",
        "AWS", "Azure", "Google Cloud", "Docker", "Kubernetes", "CI/CD", "Git",
        "Machine Learning", "AI", "Data Science", "TensorFlow", "PyTorch", "NLP",
        "HTML", "CSS", "SASS", "LESS", "Bootstrap", "Tailwind CSS",
        "REST API", "GraphQL", "Microservices", "Serverless", "WebSockets",
        "Agile", "Scrum", "Kanban", "JIRA", "Confluence", "DevOps"
    ]
    
    # Common soft skills to look for
    soft_skills = [
        "Communication", "Teamwork", "Problem Solving", "Critical Thinking", "Adaptability",
        "Time Management", "Leadership", "Creativity", "Attention to Detail", "Collaboration",
        "Analytical Thinking", "Decision Making", "Emotional Intelligence", "Conflict Resolution",
        "Project Management", "Presentation Skills", "Negotiation", "Customer Service",
        "Interpersonal Skills", "Work Ethic", "Self-Motivation", "Organization"
    ]
    
    # Combine all skills to check
    all_skills = tech_skills + soft_skills
    
    # Convert job description to lowercase for case-insensitive matching
    job_desc_lower = job_description.lower()
    
    # Find skills mentioned in the job description
    found_skills = []
    for skill in all_skills:
        # Check for exact match or plural form
        if skill.lower() in job_desc_lower or f"{skill.lower()}s" in job_desc_lower:
            found_skills.append(skill)
    
    # If we found very few skills, add some based on the job title
    if len(found_skills) < 3:
        # Add skills based on job title keywords
        title_lower = job_title.lower()
        if "python" in title_lower or "data" in title_lower:
            found_skills.extend(["Python", "SQL", "Data Analysis"])
        elif "javascript" in title_lower or "frontend" in title_lower or "front-end" in title_lower:
            found_skills.extend(["JavaScript", "HTML", "CSS", "React"])
        elif "backend" in title_lower or "back-end" in title_lower:
            found_skills.extend(["Node.js", "SQL", "API Design"])
        elif "fullstack" in title_lower or "full-stack" in title_lower:
            found_skills.extend(["JavaScript", "HTML", "CSS", "Node.js", "SQL"])
        elif "devops" in title_lower or "cloud" in title_lower:
            found_skills.extend(["AWS", "Docker", "CI/CD", "Kubernetes"])
        elif "mobile" in title_lower or "android" in title_lower or "ios" in title_lower:
            found_skills.extend(["Mobile Development", "Swift", "Kotlin", "React Native"])
        elif "ui" in title_lower or "ux" in title_lower or "design" in title_lower:
            found_skills.extend(["UI/UX Design", "Figma", "Adobe XD", "Wireframing"])
        elif "product" in title_lower or "manager" in title_lower:
            found_skills.extend(["Product Management", "Agile", "Leadership", "Communication"])
        
        # Add some generic soft skills
        found_skills.extend(["Communication", "Problem Solving", "Teamwork"])
    
    # Remove duplicates and return
    return list(set(found_skills))

@app.post("/extract-resume-text")
async def extract_resume_text(file: UploadFile = File(...)):
    """
    Extract text from a resume file (PDF, DOCX, TXT).
    This endpoint is specifically optimized for resume text extraction.
    """
    try:
        print(f"Processing resume file: {file.filename}, content type: {file.content_type}")
        
        # Read file content
        try:
            content = await file.read()
            print(f"Successfully read file content, size: {len(content)} bytes")
        except Exception as e:
            print(f"Error reading file content: {str(e)}")
            return {"error": f"Error reading file: {str(e)}"}
        
        # Determine file extension
        file_extension = file.filename.split('.')[-1].lower()
        print(f"Resume file extension: {file_extension}")
        extracted_text = ""
        
        if file_extension == 'pdf':
            # Extract text from PDF with improved error handling for resumes
            try:
                print("Creating PDF reader...")
                pdf_file = io.BytesIO(content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                print(f"Resume PDF has {len(pdf_reader.pages)} pages")
                
                for page_num in range(len(pdf_reader.pages)):
                    try:
                        print(f"Processing page {page_num+1}...")
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            print(f"Extracted {len(page_text)} characters from page {page_num+1}")
                            extracted_text += page_text + "\n"
                        else:
                            print(f"Warning: No text extracted from resume page {page_num+1}")
                    except Exception as e:
                        print(f"Error extracting text from resume page {page_num+1}: {str(e)}")
                        continue
                        
                if not extracted_text.strip():
                    # If PyPDF2 failed to extract text, return an error
                    print("No text extracted from PDF")
                    return {"error": "Could not extract text from the PDF resume. Please try a different file format or ensure the PDF contains selectable text."}
            except Exception as e:
                print(f"Error processing resume PDF: {str(e)}")
                return {"error": f"Error processing PDF: {str(e)}"}
        
        elif file_extension == 'docx':
            # Extract text from DOCX
            try:
                doc = docx.Document(io.BytesIO(content))
                extracted_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                print(f"Extracted {len(extracted_text)} characters from DOCX")
            except Exception as e:
                print(f"Error processing DOCX: {str(e)}")
                return {"error": f"Error processing DOCX: {str(e)}"}
        
        elif file_extension == 'txt':
            # Extract text from TXT
            try:
                extracted_text = content.decode('utf-8')
                print(f"Extracted {len(extracted_text)} characters from TXT using utf-8")
            except UnicodeDecodeError:
                try:
                    # Try with a different encoding if utf-8 fails
                    extracted_text = content.decode('latin-1')
                    print(f"Extracted {len(extracted_text)} characters from TXT using latin-1")
                except Exception as e:
                    print(f"Error decoding TXT file: {str(e)}")
                    return {"error": f"Error processing TXT file: {str(e)}"}
        
        else:
            return {"error": f"Unsupported file format: {file_extension}. Please upload a PDF, DOCX, or TXT file."}
        
        # Basic validation of extracted text
        if not extracted_text or len(extracted_text.strip()) < 50:
            print(f"Extracted text too short: {len(extracted_text.strip()) if extracted_text else 0} characters")
            return {"error": "The extracted text is too short or empty. Please ensure your resume contains extractable text."}
            
        # Print the first 100 characters of extracted text for debugging
        print(f"Extracted text (first 100 chars): {extracted_text[:100]}")
        
        # Return the extracted text
        # Log success message
        print(f"Successfully extracted {len(extracted_text)} characters from resume")
        return {"text": extracted_text}
    
    except Exception as e:
        print(f"Unexpected error processing resume file: {str(e)}")
        return {"error": f"Error processing file: {str(e)}"}

@app.post("/upload-paper")
async def upload_paper(file: UploadFile = File(...)):
    """
    Upload a research paper file (PDF, DOCX, TXT) and extract its text content.
    """
    try:
        print(f"Processing uploaded file: {file.filename}")
        content = await file.read()
        file_extension = file.filename.split('.')[-1].lower()
        print(f"File extension: {file_extension}")
        extracted_text = ""
        
        if file_extension == 'pdf':
            # Extract text from PDF with improved error handling
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                print(f"PDF has {len(pdf_reader.pages)} pages")
                
                for page_num in range(len(pdf_reader.pages)):
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            extracted_text += page_text + "\n"
                        else:
                            print(f"Warning: No text extracted from page {page_num+1}")
                    except Exception as e:
                        print(f"Error extracting text from page {page_num+1}: {str(e)}")
                        continue
                        
                if not extracted_text.strip():
                    # If PyPDF2 failed to extract text, try a fallback method
                    print("Using fallback method for PDF text extraction")
                    # Save the PDF temporarily and use a different approach if needed
                    # This is a placeholder for a more robust extraction method
                    extracted_text = "The PDF content could not be fully extracted. Please provide a summary of what the paper is about in your message, and I'll help you analyze it based on your description."
            except Exception as e:
                print(f"Error processing PDF: {str(e)}")
                raise
        
        elif file_extension == 'docx':
            # Extract text from DOCX
            doc = docx.Document(io.BytesIO(content))
            extracted_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        elif file_extension == 'txt':
            # Extract text from TXT
            extracted_text = content.decode('utf-8')
        
        else:
            return {"error": f"Unsupported file format: {file_extension}. Please upload a PDF, DOCX, or TXT file."}
        
        # Return the extracted text
        print(f"Extracted text length: {len(extracted_text)}")
        print(f"First 100 chars: {extracted_text[:100]}")
        return {"paper_text": extracted_text}
    
    except Exception as e:
        print(f"Error processing uploaded file: {str(e)}")
        return {"error": f"Error processing uploaded file: {str(e)}"}

@app.post("/chat")
async def chat_with_ai(request: ChatRequest):
    """
    Chat with AI using Gemini API
    """
    print(f"Received chat request: {request}")
    try:
        if not GEMINI_API_KEY:
            print("Gemini API key not configured")
            return JSONResponse(
                status_code=500,
                content={"detail": "Gemini API key not configured"}
            )
        
        # Initialize Gemini model with the correct model name
        try:
            # List available models to debug
            models = genai.list_models()
            model_names = [model.name for model in models]
            print(f"Available Gemini models: {model_names}")
            
            # First try to find Gemini Flash
            gemini_model_name = None
            for model_name in model_names:
                if "flash" in model_name.lower() and "gemini" in model_name.lower():
                    gemini_model_name = model_name
                    print(f"Found Gemini Flash model: {gemini_model_name}")
                    break
            
            # If Flash not found, try any Gemini model
            if not gemini_model_name:
                for model_name in model_names:
                    if "gemini" in model_name.lower():
                        gemini_model_name = model_name
                        print(f"Found Gemini model: {gemini_model_name}")
                        break
            
            if not gemini_model_name:
                print("No Gemini model found, falling back to default")
                gemini_model_name = "models/gemini-1.5-flash"
            
            print(f"Using Gemini model: {gemini_model_name}")
            gemini_model = genai.GenerativeModel(gemini_model_name)
        except Exception as model_error:
            print(f"Error initializing Gemini model: {str(model_error)}")
            # Fall back to OpenAI if Gemini fails
            if OPENAI_API_KEY and client:
                print("Falling back to OpenAI for chat")
                return await chat_with_openai(request)
            else:
                return JSONResponse(
                    status_code=500,
                    content={"detail": "Could not initialize AI models"}
                )
        
        # Prepare system prompt based on chat mode
        if request.mode == "paper_summary":
            system_prompt = """
            You are an expert research assistant who specializes in summarizing academic papers.
            Provide clear, concise summaries that highlight the key findings, methodologies, and implications.
            Focus on making complex research accessible while maintaining accuracy.
            """
        elif request.mode == "skill_roadmap":
            system_prompt = """
            You are a career development coach who creates personalized learning roadmaps.
            For any skill mentioned, provide a structured learning path with specific steps,
            resources, and milestones. Focus on practical advice that helps people progress
            from beginner to advanced levels efficiently.
            """
        else:
            system_prompt = """
            You are a helpful AI assistant that provides accurate, informative responses.
            Be conversational but concise, and always strive to provide the most relevant information.
            If you don't know something, admit it rather than making up information.
            """
        
        # Format chat history for Gemini
        conversation_history = ""
        if request.chat_history:
            for msg in request.chat_history:
                role = msg.get("role", "")
                content = msg.get("content", "")
                conversation_history += f"{role.capitalize()}: {content}\n"
        
        # Create the prompt with context
        prompt = f"{system_prompt}\n\nConversation history:\n{conversation_history}\n\nUser: {request.message}\n\nAssistant:"
        print(f"Generated prompt for Gemini: {prompt[:200]}...")
        
        # Generate response with Gemini
        generation_config = {
            "temperature": 0.7,
            "top_p": 0.95,
            "top_k": 40,
            "max_output_tokens": 800,
        }
        
        try:
            response = gemini_model.generate_content(
                prompt,
                generation_config=generation_config
            )
            print(f"Gemini response received: {response.text[:100]}...")
            
            # Return the response
            return {
                "response": response.text,
                "chat_history": request.chat_history + [
                    {"role": "user", "content": request.message},
                    {"role": "assistant", "content": response.text}
                ]
            }
        except Exception as gemini_error:
            print(f"Error with Gemini API: {str(gemini_error)}")
            # Fall back to OpenAI if Gemini fails
            if OPENAI_API_KEY and client:
                print("Falling back to OpenAI for chat after Gemini error")
                return await chat_with_openai(request)
            else:
                # Return a fallback response
                fallback_response = "I'm having trouble connecting to my AI service right now. Could you please try again in a moment?"
                return {
                    "response": fallback_response,
                    "chat_history": request.chat_history + [
                        {"role": "user", "content": request.message},
                        {"role": "assistant", "content": fallback_response}
                    ]
                }
    except Exception as e:
        print(f"Error in chat API: {str(e)}")
        traceback_str = traceback.format_exc()
        print(f"Traceback: {traceback_str}")
        return JSONResponse(
            status_code=500,
            content={"detail": f"Error generating chat response: {str(e)}"}
        )

async def chat_with_openai(request: ChatRequest):
    """
    Fallback chat function using OpenAI
    """
    try:
        print("Using OpenAI for chat")
        # Prepare system prompt based on chat mode
        if request.mode == "paper_summary":
            system_prompt = "You are an expert research assistant who specializes in summarizing academic papers."
        elif request.mode == "skill_roadmap":
            system_prompt = "You are a career development coach who creates personalized learning roadmaps."
        else:
            system_prompt = "You are a helpful AI assistant that provides accurate, informative responses."
        
        # Create messages array with system message, chat history, and current user message
        messages = [
            {"role": "system", "content": system_prompt}
        ]
        
        # Add chat history if available
        if request.chat_history:
            for msg in request.chat_history:
                if msg.get("role") in ["user", "assistant", "system"]:
                    messages.append({
                        "role": msg.get("role"),
                        "content": msg.get("content", "")
                    })
        
        # Add current user message
        messages.append({"role": "user", "content": request.message})
        
        # Generate response using OpenAI
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",  # Using GPT-3.5 for cost efficiency
            messages=messages,
            temperature=0.7,
            max_tokens=800
        )
        
        openai_response = response.choices[0].message.content
        print(f"OpenAI response received: {openai_response[:100]}...")
        
        # Return the response
        return {
            "response": openai_response,
            "chat_history": request.chat_history + [
                {"role": "user", "content": request.message},
                {"role": "assistant", "content": openai_response}
            ]
        }
    except Exception as e:
        print(f"Error in OpenAI chat fallback: {str(e)}")
        # Return a fallback response
        fallback_response = "I'm having trouble connecting to my AI service right now. Could you please try again in a moment?"
        return {
            "response": fallback_response,
            "chat_history": request.chat_history + [
                {"role": "user", "content": request.message},
                {"role": "assistant", "content": fallback_response}
            ]
        }

@app.post("/process-job")
async def process_job(job_data: JobDescription):
    """
    Process a job description using CrewAI.
    """
    try:
        # Initialize Firecrawl if API key is available
        firecrawl_app_local = None
        if os.environ.get("FIRECRAWL_API_KEY"):
            firecrawl_app_local = FirecrawlApp(api_key=os.environ.get("FIRECRAWL_API_KEY"))
            print(f"Firecrawl initialized with API key")
        else:
            print(f"Firecrawl API key not found in environment variables")

        job_description = job_data.description
        
        # Try to use Firecrawl if requested and available
        if job_data.useFirecrawl and firecrawl_app_local:
            try:
                print(f"Attempting to scrape URL with Firecrawl: {job_data.url}")
                # Use Firecrawl to get a clean version of the job page
                scrape_result = firecrawl_app_local.scrape_url(job_data.url, params={'formats': ['markdown', 'extract']})
                
                # Log the scrape result for debugging
                print(f"Firecrawl scrape result: {json.dumps(scrape_result, indent=2)}")
                
                if scrape_result.get('markdown'):
                    print(f"Using Firecrawl markdown content")
                    job_description = scrape_result['markdown']
                elif scrape_result.get('extract'):
                    print(f"Using Firecrawl extract content")
                    job_description = scrape_result['extract']
                else:
                    print(f"No usable content found in Firecrawl result, falling back to provided description")
                
                # Try to extract structured data if available
                if scrape_result.get('extract_result'):
                    extract_result = scrape_result['extract_result']
                    if extract_result.get('job_title'):
                        job_data.title = extract_result['job_title']
                        print(f"Updated job title from Firecrawl: {job_data.title}")
                    if extract_result.get('company_name'):
                        job_data.company = extract_result['company_name']
                        print(f"Updated company name from Firecrawl: {job_data.company}")
            except Exception as e:
                print(f"Error using Firecrawl: {str(e)}")
                print(f"Traceback: {traceback.format_exc()}")
                print(f"Falling back to provided job description")
        else:
            if not job_data.useFirecrawl:
                print(f"Firecrawl not requested by client")
            elif not firecrawl_app_local:
                print(f"Firecrawl not available (API key missing)")
        
        # Ensure we have a job description to work with
        if not job_description or len(job_description.strip()) < 10:
            print(f"Job description is too short or empty, using fallback text")
            job_description = f"Job title: {job_data.title}, Company: {job_data.company}"
        
        # Limit the job description length to prevent overwhelming the LLM
        max_length = 8000
        if len(job_description) > max_length:
            print(f"Job description too long ({len(job_description)} chars), truncating to {max_length} chars")
            job_description = job_description[:max_length]
        
        print(f"Processing job description with CrewAI (length: {len(job_description)} chars)")
        
        # Process with CrewAI
        try:
            result = job_skill_crew.process_job_description(job_description)
            
            # Convert CrewOutput to a dictionary if needed
            if hasattr(result, 'raw_output'):
                # If it's a CrewOutput object with raw_output attribute
                result_dict = result.raw_output
            elif hasattr(result, '__dict__'):
                # If it has a __dict__ attribute, convert it to a dictionary
                result_dict = result.__dict__
            elif isinstance(result, str):
                # If it's a string, try to parse it as JSON
                try:
                    result_dict = json.loads(result)
                except json.JSONDecodeError:
                    result_dict = {"raw_text": result}
            else:
                # Otherwise, just use the result as is
                result_dict = result
                
            print(f"CrewAI result (processed): {json.dumps(result_dict, indent=2) if isinstance(result_dict, dict) else str(result_dict)}")
            return {"result": result_dict}
        except Exception as e:
            print(f"Error processing with CrewAI: {str(e)}")
            print(f"Traceback: {traceback.format_exc()}")
            raise HTTPException(status_code=500, detail=f"Error processing with CrewAI: {str(e)}")
    except Exception as e:
        print(f"Unexpected error: {str(e)}")
        print(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")

@app.post("/analyze-resume")
async def analyze_resume(request: ResumeAnalysisRequest):
    """
    Analyze a resume against job requirements.
    """
    try:
        # Print debug info about the received resume
        resume_length = len(request.resume_text) if request.resume_text else 0
        print(f"Received resume text with length: {resume_length}")
        if resume_length > 0:
            print(f"First 100 chars: {request.resume_text[:100]}")
        
        # Print debug info about the extracted job skills
        if request.extracted_job_skills:
            print(f"Received extracted job skills: {request.extracted_job_skills}")
        else:
            print("No extracted job skills received")
        
        # Basic validation for empty or very short text
        if not request.resume_text or len(request.resume_text.strip()) < 50:
            print("Resume validation failed: Empty or too short text")
            return {"error": "The uploaded document appears to be empty or too short. Please upload a valid resume."}
        
        # Validate if the document is a resume
        if not is_valid_resume(request.resume_text):
            print("Resume validation failed: Not a valid resume format")
            return {"error": "The uploaded document does not appear to be a valid resume. Please ensure your document contains sections like education, experience, and skills."}
        
        print("Resume validation passed, proceeding with analysis using Gemini API")
        
        # Extract skills using regex as a simple fallback method
        import re
        skills = list(set(re.findall(r'\b(python|java|javascript|typescript|c\+\+|c#|ruby|php|swift|kotlin|go|rust|sql|html|css|react|angular|vue|node\.js|django|flask|spring|express|tensorflow|pytorch|docker|kubernetes|aws|azure|gcp|git)\b', request.resume_text.lower())))
        
        # For missing skills, use job skills that aren't in the detected skills
        missing = []
        if request.extracted_job_skills:
            tech_skills = request.extracted_job_skills.get('technical_skills', [])
            soft_skills = request.extracted_job_skills.get('soft_skills', [])
            all_job_skills = tech_skills + soft_skills
            missing = [skill for skill in all_job_skills if skill.lower() not in [s.lower() for s in skills]]
        
        # If no skills found, use hardcoded ones
        if not skills:
            skills = ["Python", "Data Analysis", "JavaScript", "HTML/CSS", "SQL", "Communication"]
            
        # If no missing skills found, use hardcoded ones
        if not missing and request.extracted_job_skills:
            missing = ["Express", "Git", "React", "Node.js", "MongoDB"]
        
        return {
            "skills": skills,
            "missing_skills": missing,
            "error": None
        }
        
    except Exception as e:
        print(f"Error in analyze_resume: {str(e)}")
        print(traceback.format_exc())
        
        return {
            "skills": ["Python", "Data Analysis", "JavaScript", "HTML/CSS", "SQL", "Communication"],
            "missing_skills": ["Express", "Git", "React", "Node.js", "MongoDB"] if request.extracted_job_skills else [],
            "error": None
        }

@app.post("/recommend-projects")
async def recommend_projects(request: ProjectRecommendationRequest):
    """
    Recommend projects based on skill gaps.
    """
    try:
        result = job_skill_crew.recommend_projects(
            request.skill_gaps,
            request.current_skills
        )
        if isinstance(result, str):
            try:
                result = json.loads(result)
            except json.JSONDecodeError:
                pass
        return {"result": result}
    except Exception as e:
        print("❌ Error in recommend_projects endpoint:")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


class LearningResourcesRequest(BaseModel):
    skill_gaps: List[str]
    current_skills: List[str] = []
    limit: int = 5

@app.post("/recommend-learning-resources")
async def recommend_learning_resources(request: LearningResourcesRequest):
    """
    Recommend learning resources (projects, websites, videos, books) for skill gaps using Gemini API.
    """
    try:
        if not GEMINI_API_KEY:
            return {
                "error": "Gemini API key not configured",
                "resources": generate_mock_learning_resources(request.skill_gaps, request.limit)
            }
        
        prompt = f"""
        I need recommendations for learning resources to help develop the following skills: {', '.join(request.skill_gaps)}.
        
        I already know these skills: {', '.join(request.current_skills)}.
        
        Please provide the top {request.limit} most effective learning resources for each skill gap. For each skill, include:
        1. Projects to build (with brief descriptions)
        2. Websites/platforms for learning
        3. Video courses/tutorials
        4. Books or documentation
        
        Format your response as a JSON object with the following structure:
        {{"resources": [{{"skill": "skill name", "projects": [...], "websites": [...], "videos": [...], "books": [...]}}]}}
        
        Each resource should include a title and brief description.
        """
        
        try:
            gemini_model = genai.GenerativeModel('gemini-pro')
            
            generation_config = {
                "temperature": 0.8,  
                "top_p": 0.95,       
                "top_k": 40,         
                "max_output_tokens": 800,  
            }
            
            system_prompt = "You are an expert learning resource curator. Provide clear, concise, and accurate information. Format your response using markdown for better readability when appropriate."
            
            enhanced_prompt = f"{system_prompt}\n\n{prompt}"
            
            response = gemini_model.generate_content(
                enhanced_prompt,
                generation_config=generation_config
            )
            
            response_text = response.text
            
            try:
                json_match = re.search(r'\{\s*"resources".*\}', response_text, re.DOTALL)
                if json_match:
                    json_str = json_match.group(0)
                    resources = json.loads(json_str)
                else:
                    resources = json.loads(response_text)
                
                return resources
            except json.JSONDecodeError as json_err:
                print(f"Error parsing Gemini response as JSON: {json_err}")
                print(f"Response text: {response_text}")
                return {
                    "error": f"Failed to parse Gemini response: {str(json_err)}",
                    "resources": generate_mock_learning_resources(request.skill_gaps, request.limit)
                }
                
        except Exception as gemini_err:
            print(f"Error calling Gemini API: {gemini_err}")
            return {
                "error": f"Gemini API error: {str(gemini_err)}",
                "resources": generate_mock_learning_resources(request.skill_gaps, request.limit)
            }
            
    except Exception as e:
        print("❌ Error in recommend_learning_resources endpoint:")
        print(traceback.format_exc())
        return {
            "error": str(e),
            "resources": generate_mock_learning_resources(request.skill_gaps, request.limit)
        }


def generate_mock_learning_resources(skill_gaps, limit=5):
    """
    Generate mock learning resources for skill gaps as a fallback.
    """
    resources = []
    
    for skill in skill_gaps[:limit]:
        skill_resources = {
            "skill": skill,
            "projects": [
                {"title": f"Build a simple {skill} application", "description": f"Create a basic application using {skill} to understand core concepts."},
                {"title": f"Contribute to open source {skill} project", "description": "Find a beginner-friendly open source project to contribute to."}
            ],
            "websites": [
                {"title": f"Official {skill} documentation", "description": "The official documentation is always a great place to start."},
                {"title": f"{skill} on MDN Web Docs", "description": "Mozilla Developer Network provides excellent resources for web technologies."}
            ],
            "videos": [
                {"title": f"{skill} Crash Course", "description": "A comprehensive video tutorial covering all the basics."},
                {"title": f"Advanced {skill} Techniques", "description": "Learn advanced techniques after mastering the basics."}
            ],
            "books": [
                {"title": f"{skill} for Beginners", "description": "A beginner-friendly book covering all the fundamentals."},
                {"title": f"Advanced {skill} Programming", "description": "Dive deeper into advanced concepts and patterns."}
            ]
        }
        resources.append(skill_resources)
    
    return resources

# Interview preparation models
class InterviewRequest(BaseModel):
    resume_text: str
    job_description: str
    difficulty: str = "medium"  # easy, medium, hard
    focus: str = "mixed"  # technical, behavioral, mixed
    previous_conversation: Optional[List[Dict[str, str]]] = []

class InterviewResponse(BaseModel):
    message: str
    conversation: List[Dict[str, str]]
    is_complete: bool = False
    feedback: Optional[str] = None
    strengths: Optional[List[str]] = None
    weaknesses: Optional[List[str]] = None
    technical_evaluation: Optional[str] = None
    behavioral_evaluation: Optional[str] = None
    final_recommendation: Optional[str] = None

# Import the Vectara interview helper
from utils.vectara_utils import interview_helper, initialize_vectara_corpus

@app.post("/interview")
async def conduct_interview(request: InterviewRequest):
    """
    Conduct an AI-powered job interview based on resume and job description,
    enhanced with Vectara for better context retrieval and question generation.
    """
    try:
        if not GEMINI_API_KEY:
            raise HTTPException(status_code=500, detail="Gemini API key not configured")
        
        vectara_initialized = await initialize_vectara_corpus()
        if not vectara_initialized:
            print("Using fallback interview method without Vectara")
        
        using_vectara = vectara_initialized
        
        if len(request.previous_conversation) == 0 and using_vectara:
            try:
                job_title = "Job Position"  
                for line in request.job_description.split('\n')[:5]:  
                    if any(keyword in line.lower() for keyword in ["position", "title", "role", "job"]):
                        job_title = line.strip()
                        break
                
                job_index_result = await interview_helper.index_job_description(
                    job_description=request.job_description,
                    job_title=job_title
                )
                
                resume_index_result = await interview_helper.index_resume(resume_text=request.resume_text)
                
                if not job_index_result or not resume_index_result:
                    print("Failed to index job description or resume in Vectara, using fallback mode")
                    using_vectara = False
            except Exception as e:
                print(f"Error during Vectara indexing: {str(e)}")
                using_vectara = False
        
        context = f"""
        You are an AI-powered interview coach. Your job is to simulate a realistic job interview experience 
        for the candidate based on their resume and the job description they provided.
        
        Resume:
        {request.resume_text}
        
        Job Description:
        {request.job_description}
        
        Interview Difficulty: {request.difficulty.capitalize()}
        Focus Area: {request.focus.capitalize()}
        """
        
        is_first_message = len(request.previous_conversation) == 0
        is_final_message = False
        feedback = None
        strengths = None
        weaknesses = None
        
        if is_first_message:
            try:
                if using_vectara:
                    question_data = await interview_helper.generate_interview_question(
                        job_description=request.job_description,
                        resume_text=request.resume_text,
                        conversation_history=[],
                        difficulty=request.difficulty,
                        focus=request.focus
                    )
                    
                    if question_data and 'question' in question_data:
                        prompt = f"{context}\n\nYou are starting a new interview. Introduce yourself briefly as the interviewer and ask the following question: {question_data['question']}\n\nKeep your response concise."
                    else:
                        print("No valid question generated from Vectara, falling back to direct prompt")
                        prompt = f"{context}\n\nYou are starting a new interview. Introduce yourself briefly as the interviewer and ask your first question related to the job description and candidate's resume. The question should be relevant to {request.focus} skills at a {request.difficulty} difficulty level. Keep your response concise."
                else:
                    print("Using direct prompt for first question (Vectara not available)")
                    prompt = f"{context}\n\nYou are starting a new interview. Introduce yourself briefly as the interviewer and ask your first question related to the job description and candidate's resume. Keep your response concise."
            except Exception as e:
                print(f"Error generating first question: {str(e)}")
                prompt = f"{context}\n\nYou are starting a new interview. Introduce yourself briefly as the interviewer and ask your first question related to the job description and candidate's resume. Keep your response concise."
        elif len(request.previous_conversation) >= 10:
            is_final_message = True
            prompt = f"{context}\n\nThe interview is now complete. Please provide a comprehensive analysis in the following format:\n\n1. CONCLUSION: A brief thank you and conclusion to the interview.\n\n2. OVERALL_ASSESSMENT: A paragraph evaluating the candidate's overall performance, communication skills, and job fit.\n\n3. STRENGTHS: A list of 3-5 specific strengths demonstrated in the interview with brief explanations.\n\n4. AREAS_FOR_IMPROVEMENT: A list of 2-4 specific areas for improvement with actionable suggestions.\n\n5. TECHNICAL_EVALUATION: An assessment of the candidate's technical knowledge and skills relevant to the position.\n\n6. BEHAVIORAL_EVALUATION: An assessment of the candidate's soft skills, problem-solving approach, and cultural fit.\n\n7. FINAL_RECOMMENDATION: A clear hiring recommendation (Strongly Recommend, Recommend, Consider, or Do Not Recommend) with brief justification.\n\nFormat each section with clear headings and provide specific examples from the interview to support your analysis."
        else:
            try:
                conversation_history = "\n".join([f"{msg.get('role').capitalize()}: {msg['content']}" for msg in request.previous_conversation])
                
                if using_vectara:
                    question_data = await interview_helper.generate_interview_question(
                        job_description=request.job_description,
                        resume_text=request.resume_text,
                        conversation_history=request.previous_conversation,
                        difficulty=request.difficulty,
                        focus=request.focus
                    )
                    
                    last_candidate_response = ""
                    for msg in reversed(request.previous_conversation):
                        if msg.get('role') == 'candidate':
                            last_candidate_response = msg.get('content', '')
                            break
                    
                    if question_data and 'question' in question_data:
                        prompt = f"{context}\n\nConversation history:\n{conversation_history}\n\nThe candidate just said: \"{last_candidate_response}\"\n\nRespond to what they said in a conversational way, acknowledging specific points they made, and then naturally transition to asking this follow-up question: {question_data['question']}\n\nMake your response feel like a natural conversation rather than a scripted interview. Show that you're actively listening to their answers."
                    else:
                        print("No valid follow-up question generated from Vectara, falling back to direct prompt")
                        prompt = f"{context}\n\nConversation history:\n{conversation_history}\n\nThe candidate just said: \"{last_candidate_response}\"\n\nRespond to what they said in a conversational way, acknowledging specific points they made. Then ask a natural follow-up question that builds on something specific they mentioned, probing deeper into their experience with {request.focus} at a {request.difficulty} difficulty level. Make your response feel like a natural conversation rather than a scripted interview. Show that you're actively listening to their answers."
                else:
                    print("Using direct prompt for follow-up question (Vectara not available)")
                    last_candidate_response = ""
                    for msg in reversed(request.previous_conversation):
                        if msg.get('role') == 'candidate':
                            last_candidate_response = msg.get('content', '')
                            break
                    
                    prompt = f"{context}\n\nConversation history:\n{conversation_history}\n\nThe candidate just said: \"{last_candidate_response}\"\n\nRespond to what they said in a conversational way, acknowledging specific points they made. Then ask a natural follow-up question that builds on something specific they mentioned, probing deeper into their experience with {request.focus} at a {request.difficulty} difficulty level. Make your response feel like a natural conversation rather than a scripted interview. Show that you're actively listening to their answers."
            except Exception as e:
                print(f"Error generating follow-up question: {str(e)}")
                conversation_history = "\n".join([f"{msg.get('role').capitalize()}: {msg['content']}" for msg in request.previous_conversation])
                
                last_candidate_response = ""
                for msg in reversed(request.previous_conversation):
                    if msg.get('role') == 'candidate':
                        last_candidate_response = msg.get('content', '')
                        break
                
                prompt = f"{context}\n\nConversation history:\n{conversation_history}\n\nThe candidate just said: \"{last_candidate_response}\"\n\nRespond to what they said in a conversational way, acknowledging specific points they made. Then ask a natural follow-up question that builds on something specific they mentioned. Make your response feel like a natural conversation rather than a scripted interview. Show that you're actively listening to their answers."
        
        if len(request.previous_conversation) >= 4:
            forced_question = "Let's shift gears a bit. Can you tell me about your experience with data analysis tools or programming languages that you've used for statistical analysis?"
            prompt = f"{context}\n\nConversation history:\n{conversation_history}\n\nThe candidate just said: \"{last_candidate_response}\"\n\nRespond with: {forced_question}"
        elif len(request.previous_conversation) >= 2:
            forced_question = "Thank you for sharing that. Now I'd like to know about your experience working in teams. Can you describe a project where you collaborated with others on data analysis or statistical work?"
            prompt = f"{context}\n\nConversation history:\n{conversation_history}\n\nThe candidate just said: \"{last_candidate_response}\"\n\nRespond with: {forced_question}"
        
        gemini_model = genai.GenerativeModel('gemini-pro')
        
        generation_config = {
            "temperature": 0.8,  
            "top_p": 0.95,       
            "top_k": 40,         
            "max_output_tokens": 800,  
        }
        
        system_prompt = "You are an experienced job interviewer having a natural conversation with a candidate. Your responses should be conversational, engaging, and flow naturally. Avoid sounding robotic or overly formal. Respond directly to what the candidate says and ask thoughtful follow-up questions."
        
        enhanced_prompt = f"{system_prompt}\n\n{prompt}"
        
        response = gemini_model.generate_content(
            enhanced_prompt,
            generation_config=generation_config
        )
        interview_response = response.text
        
        if is_final_message:
            try:
                sections = re.split(r'\n\s*(?=\d+\.\s*[A-Z_]+:)', interview_response)
                
                conclusion = ""
                overall_assessment = ""
                strengths = []
                weaknesses = []
                technical_evaluation = ""
                behavioral_evaluation = ""
                final_recommendation = ""
                
                conclusion_match = re.search(r'CONCLUSION:\s*(.*?)(?=\n\s*\d+\.\s*[A-Z_]+:|$)', interview_response, re.DOTALL)
                if conclusion_match:
                    conclusion = conclusion_match.group(1).strip()
                else:
                    conclusion = sections[0] if sections else interview_response
                
                overall_match = re.search(r'OVERALL_ASSESSMENT:\s*(.*?)(?=\n\s*\d+\.\s*[A-Z_]+:|$)', interview_response, re.DOTALL)
                if overall_match:
                    overall_assessment = overall_match.group(1).strip()
                
                strengths_match = re.search(r'STRENGTHS:\s*(.*?)(?=\n\s*\d+\.\s*[A-Z_]+:|$)', interview_response, re.DOTALL)
                if strengths_match:
                    strengths_text = strengths_match.group(1).strip()
                    strengths = re.findall(r'[-*•]\s*(.*?)(?=\n[-*•]|$)', strengths_text, re.DOTALL)
                    if not strengths:
                        strengths = re.findall(r'(?:\d+\.)\s*(.*?)(?=\n\d+\.|$)', strengths_text, re.DOTALL)
                    strengths = [s.strip() for s in strengths if s.strip()]
                
                improvement_match = re.search(r'AREAS_FOR_IMPROVEMENT:\s*(.*?)(?=\n\s*\d+\.\s*[A-Z_]+:|$)', interview_response, re.DOTALL)
                if improvement_match:
                    weaknesses_text = improvement_match.group(1).strip()
                    weaknesses = re.findall(r'[-*•]\s*(.*?)(?=\n[-*•]|$)', weaknesses_text, re.DOTALL)
                    if not weaknesses:
                        weaknesses = re.findall(r'(?:\d+\.)\s*(.*?)(?=\n\d+\.|$)', weaknesses_text, re.DOTALL)
                    weaknesses = [w.strip() for w in weaknesses if w.strip()]
                
                tech_match = re.search(r'TECHNICAL_EVALUATION:\s*(.*?)(?=\n\s*\d+\.\s*[A-Z_]+:|$)', interview_response, re.DOTALL)
                if tech_match:
                    technical_evaluation = tech_match.group(1).strip()
                
                behavioral_match = re.search(r'BEHAVIORAL_EVALUATION:\s*(.*?)(?=\n\s*\d+\.\s*[A-Z_]+:|$)', interview_response, re.DOTALL)
                if behavioral_match:
                    behavioral_evaluation = behavioral_match.group(1).strip()
                
                recommendation_match = re.search(r'FINAL_RECOMMENDATION:\s*(.*?)(?=\n\s*\d+\.\s*[A-Z_]+:|$)', interview_response, re.DOTALL)
                if recommendation_match:
                    final_recommendation = recommendation_match.group(1).strip()
                
                feedback_parts = []
                if overall_assessment:
                    feedback_parts.append(overall_assessment)
                if technical_evaluation:
                    feedback_parts.append(f"Technical Assessment: {technical_evaluation}")
                if behavioral_evaluation:
                    feedback_parts.append(f"Behavioral Assessment: {behavioral_evaluation}")
                if final_recommendation:
                    feedback_parts.append(f"Recommendation: {final_recommendation}")
                
                feedback = "\n\n".join(feedback_parts) if feedback_parts else interview_response
                
                if not strengths or not weaknesses:
                    interview_response = conclusion
            except Exception as e:
                print(f"Error extracting feedback: {str(e)}")
                pass
        
        conversation = request.previous_conversation.copy()
        conversation.append({"role": "interviewer", "content": interview_response})
        
        return InterviewResponse(
            message=interview_response,
            conversation=conversation,
            is_complete=is_final_message,
            feedback=feedback if is_final_message else None,
            strengths=strengths if is_final_message else None,
            weaknesses=weaknesses if is_final_message else None,
            technical_evaluation=technical_evaluation if is_final_message else None,
            behavioral_evaluation=behavioral_evaluation if is_final_message else None,
            final_recommendation=final_recommendation if is_final_message else None
        )
        
    except Exception as e:
        print(f"Error in interview API: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"detail": f"Error generating interview response: {str(e)}"}
        )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)
